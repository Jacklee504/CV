"""Build a fuller recruiter-facing CV from content/cv.json."""

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

import build_resume as resume


CONTENT = Path("content/cv.json")
OUT = Path("output/Jack_Lee_CV.docx")


def add_body(document, text, *, after=2.5, size=9.5, line=1.08):
    paragraph = document.add_paragraph(style="Resume Body")
    resume.set_paragraph_spacing(paragraph, after=after, line=line)
    run = paragraph.add_run(text)
    resume.set_font(run, size=size, color=resume.INK)
    return paragraph


def configure_cv_styles(document):
    """Apply the roomier CV scale without changing the one-page resume styles."""
    section_style = document.styles["Resume Section"]
    section_style.font.size = resume.Pt(10.9)
    section_style.paragraph_format.space_before = resume.Pt(13.5)
    section_style.paragraph_format.space_after = resume.Pt(5.5)


def add_project(document, project, num_id):
    detail = f"{project['context']} | {project['technologies']}"
    paragraphs = [resume.add_entry_header(
        document,
        project["title"],
        detail,
        detail_italic=True,
        title_url=project.get("url"),
    )]
    for bullet in project["bullets"]:
        paragraph = document.add_paragraph(style="Resume Body")
        pPr = paragraph._p.get_or_add_pPr()
        numPr = resume.OxmlElement("w:numPr")
        ilvl = resume.OxmlElement("w:ilvl")
        ilvl.set(resume.qn("w:val"), "0")
        num_id_element = resume.OxmlElement("w:numId")
        num_id_element.set(resume.qn("w:val"), str(num_id))
        numPr.append(ilvl)
        numPr.append(num_id_element)
        pPr.append(numPr)
        run = paragraph.add_run(bullet)
        resume.set_font(run, size=resume.BODY_SIZE, color=resume.INK)
        resume.set_paragraph_spacing(paragraph, after=3.2, line=1.07)
        resume.set_keep(paragraph, keep_together=True)
        paragraphs.append(paragraph)

    # Keep each project as one block, while allowing completed projects to fill page one.
    for paragraph in paragraphs[:-1]:
        resume.set_keep(paragraph, keep_with_next=True, keep_together=True)


def add_cv_skills(document, skills):
    paragraph = document.add_paragraph(style="Resume Body")
    resume.set_paragraph_spacing(paragraph, after=2.5, line=1.15)
    for index, (label, value) in enumerate(skills):
        run = paragraph.add_run(f"{label}: ")
        resume.set_font(run, size=10.15, bold=True, color=resume.INK)
        run = paragraph.add_run(value)
        resume.set_font(run, size=10.15, color=resume.INK)
        if index < len(skills) - 1:
            paragraph.add_run("\n")
    return paragraph


def add_education(document, education):
    resume.add_entry_header(
        document,
        education["degree"],
        None,
        education["dates"],
    )
    add_body(document, education["institution"], after=4.5, size=10.15, line=1.12)
    for detail in education["details"]:
        add_body(document, detail, after=4.5, size=10.15, line=1.12)

    secondary = education["secondary"]
    resume.add_entry_header(
        document,
        secondary["qualification"],
        secondary["institution"],
        secondary["dates"],
    )
    add_body(document, secondary["details"], after=4.5, size=10.15, line=1.12)


def add_page_number(paragraph):
    run = paragraph.add_run("Page ")
    resume.set_font(run, size=8.5, color=resume.MUTED)
    field = resume.OxmlElement("w:fldSimple")
    field.set(resume.qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def main():
    content = json.loads(CONTENT.read_text(encoding="utf-8"))
    OUT.parent.mkdir(parents=True, exist_ok=True)

    # A CV can breathe more than a one-page resume while retaining the same visual system.
    resume.BODY_SIZE = 9.75
    resume.ENTRY_SIZE = 10.2
    resume.DATE_SIZE = 9.55

    document = Document()
    resume.configure_document(document)
    configure_cv_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    resume.set_paragraph_spacing(footer, before=0, after=0, line=1.0)
    add_page_number(footer)

    num_id = resume.add_custom_bullet_numbering(document)
    resume.add_header(document, content)

    resume.add_section(document, "Profile")
    add_body(document, content["profile"], after=5, size=9.9)

    resume.add_section(document, "Professional Experience")
    for role in content["experience"]:
        resume.add_role(document, role, num_id)

    resume.add_section(document, "Selected Projects")
    for project in content["projects"]:
        add_project(document, project, num_id)

    resume.add_section(document, "Technical Skills")
    add_cv_skills(document, content["skills"])

    resume.add_section(document, "Education")
    add_education(document, content["education"])

    resume.add_section(document, "Professional Learning")
    add_body(document, content["professional_learning"], after=6, size=10.15, line=1.12)

    resume.add_section(document, "Activities & Interests")
    add_body(document, content["activities"], after=0, size=10.15, line=1.12)

    document.core_properties.title = f"{content['name']} – Curriculum Vitae"
    document.core_properties.author = content["name"]
    document.core_properties.subject = "Curriculum Vitae"
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
