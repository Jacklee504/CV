"""Build the recruiter-facing resume from the editable content/resume.json source."""

import json
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


CONTENT = Path("content/resume.json")
OUT = Path("output/Jack_Lee_Software_Engineering_Resume.docx")
INK = "111827"
MUTED = "4B5563"
RULE = "1F4E79"
BODY_SIZE = 9.2
ENTRY_SIZE = 9.7
DATE_SIZE = 9.2


def set_font(run, name="Arial", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_keep(paragraph, keep_with_next=False, keep_together=False):
    pPr = paragraph._p.get_or_add_pPr()
    if keep_with_next:
        pPr.append(OxmlElement("w:keepNext"))
    if keep_together:
        pPr.append(OxmlElement("w:keepLines"))


def add_hyperlink(paragraph, text, url, color=MUTED, size=9, bold=False):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(fonts)
    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), color)
    rPr.append(color_element)
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(round(size * 2)))
    rPr.append(font_size)
    if bold:
        rPr.append(OxmlElement("w:b"))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    run.append(rPr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bottom_border(paragraph, color=RULE, size="8", space="4"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_custom_bullet_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_id = 73
    num_id = 73
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    pPr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    pPr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    pPr.append(ind)
    level.append(pPr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(document, num_id, text):
    paragraph = document.add_paragraph(style="Resume Body")
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    run = paragraph.add_run(text)
    set_font(run, size=BODY_SIZE, color=INK)
    set_paragraph_spacing(paragraph, after=3.2, line=1.07)
    set_keep(paragraph, keep_together=True)


def add_section(document, title):
    paragraph = document.add_paragraph(style="Resume Section")
    paragraph.add_run(title.upper())
    return paragraph


def add_entry_header(document, title, detail, dates=None, detail_italic=False, title_url=None):
    """Use one shared title/detail/date treatment for roles, projects, and education."""
    paragraph = document.add_paragraph(style="Resume Entry")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.48), WD_TAB_ALIGNMENT.RIGHT)
    if title_url:
        add_hyperlink(paragraph, title, title_url, color=INK, size=ENTRY_SIZE, bold=True)
    else:
        run = paragraph.add_run(title)
        set_font(run, size=ENTRY_SIZE, bold=True, color=INK)
    if detail:
        run = paragraph.add_run(f" | {detail}")
        set_font(run, size=ENTRY_SIZE, color=MUTED, italic=detail_italic)
    if dates:
        paragraph.add_run("\t")
        run = paragraph.add_run(dates)
        # All right-aligned dates intentionally use this single shared format.
        set_font(run, size=DATE_SIZE, bold=True, color=MUTED)
    set_paragraph_spacing(paragraph, before=5, after=1.6, line=1.0)
    set_keep(paragraph, keep_with_next=True, keep_together=True)
    return paragraph


def add_role(document, role, num_id):
    detail = f"{role['organisation']}, {role['location']}"
    add_entry_header(document, role["title"], detail, role["dates"])
    for item in role["bullets"]:
        add_bullet(document, num_id, item)


def add_project(document, project, num_id):
    add_entry_header(
        document,
        project["title"],
        project["technologies"],
        detail_italic=True,
        title_url=project.get("url"),
    )
    for item in project["bullets"]:
        add_bullet(document, num_id, item)


def add_skills(document, skills):
    paragraph = document.add_paragraph(style="Resume Body")
    set_paragraph_spacing(paragraph, after=0, line=1.05)
    for index, (label, value) in enumerate(skills):
        run = paragraph.add_run(f"{label}: ")
        set_font(run, size=9.15, bold=True, color=INK)
        run = paragraph.add_run(value)
        set_font(run, size=9.15, color=INK)
        if index < len(skills) - 1:
            paragraph.add_run("\n")


def add_secondary_education(document, education):
    """Separate school qualifications from university coursework without adding a second timeline row."""
    paragraph = document.add_paragraph(style="Resume Body")
    set_paragraph_spacing(paragraph, before=3, after=1, line=1.03)
    run = paragraph.add_run(education["qualification"])
    set_font(run, size=9.15, bold=True, color=INK)
    run = paragraph.add_run(f" | {education['details']}")
    set_font(run, size=9.15, color=INK)


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = RGBColor.from_string(INK)

    for name in ("Resume Body", "Resume Entry", "Resume Section"):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal

    section_style = styles["Resume Section"]
    section_style.font.name = "Arial"
    section_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    section_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    section_style.font.size = Pt(10.2)
    section_style.font.bold = True
    section_style.font.color.rgb = RGBColor.from_string(RULE)
    section_style.paragraph_format.space_before = Pt(10.5)
    section_style.paragraph_format.space_after = Pt(4)


def add_header(document, content):
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name, after=1, line=1.0)
    run = name.add_run(content["name"].upper())
    set_font(run, size=19, bold=True, color=INK)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact, after=5, line=1.0)
    add_bottom_border(contact)
    run = contact.add_run(f"{content['location']}  |  ")
    set_font(run, size=8.8, color=MUTED)
    run = contact.add_run("Email: ")
    set_font(run, size=8.8, color=MUTED)
    add_hyperlink(contact, content["email"], f"mailto:{content['email']}")
    for link in content["links"]:
        run = contact.add_run("  |  ")
        set_font(run, size=8.8, color=MUTED)
        run = contact.add_run(f"{link['label']}: ")
        set_font(run, size=8.8, color=MUTED)
        add_hyperlink(contact, link["username"], link["url"])


def main():
    content = json.loads(CONTENT.read_text(encoding="utf-8"))
    OUT.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_document(document)
    num_id = add_custom_bullet_numbering(document)
    add_header(document, content)

    summary = document.add_paragraph(style="Resume Body")
    set_paragraph_spacing(summary, after=2, line=1.07)
    run = summary.add_run(content["summary"])
    set_font(run, size=9.4, color=INK)

    add_section(document, "Experience")
    for role in content["experience"]:
        add_role(document, role, num_id)

    add_section(document, "Projects")
    for project in content["projects"]:
        add_project(document, project, num_id)

    add_section(document, "Technical Skills")
    add_skills(document, content["skills"])

    add_section(document, "Education")
    education = content["education"]
    add_entry_header(document, education["degree"], None, education["dates"])
    details = document.add_paragraph(style="Resume Body")
    run = details.add_run(education["details"])
    set_font(run, size=9.15, color=INK)
    set_paragraph_spacing(details, after=1, line=1.03)
    add_secondary_education(document, education["secondary"])

    document.core_properties.title = f"{content['name']} – Resume"
    document.core_properties.author = content["name"]
    document.core_properties.subject = "Software Engineering Resume"
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
